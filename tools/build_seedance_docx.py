"""Build docs/seedance-vial-prompts.docx from the markdown source.

Renders the markdown, then appends an appendix holding the eight finished,
paste-ready prompts read from assets/video/_seed/prompt-N.txt — so the Word
document is self-contained and nobody has to splice a label line by hand.

    python tools/emit_seedance_prompts.py     # writes the 8 prompt files
    python tools/build_seedance_docx.py       # writes the .docx

Needs: pip install python-docx
"""
import re
import pathlib

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

ROOT = pathlib.Path(__file__).resolve().parent.parent
SRC = ROOT / 'docs' / 'seedance-vial-prompts.md'
DST = ROOT / 'docs' / 'seedance-vial-prompts.docx'
SEED = ROOT / 'assets' / 'video' / '_seed'

PRODUCTS = {
    1: 'Retatrutide', 3: 'GHK-Cu',
    4: 'Tesamorelin / Ipamorelin Blend', 5: 'MOTS-C',
    6: 'BPC-157 / TB-500 Blend', 7: 'KLOW Blend', 8: 'NAD+',
}

CODE_FONT = 'Consolas'
BODY_FONT = 'Calibri'
CODE_BG = 'F2F1F6'
RULE_COL = 'C9C7D0'
INK = RGBColor(0x1A, 0x14, 0x2E)


def shade(el, fill):
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:fill'), fill)
    el.append(s)


def borders(p, fill):
    pbdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '6')
        b.set(qn('w:space'), '6')
        b.set(qn('w:color'), fill)
        pbdr.append(b)
    p._p.get_or_add_pPr().append(pbdr)


def hrule(doc):
    p = doc.add_paragraph()
    pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), RULE_COL)
    pbdr.append(b)
    p._p.get_or_add_pPr().append(pbdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)


def code_block(doc, text, size=8.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.12)
    shade(p._p.get_or_add_pPr(), CODE_BG)
    borders(p, RULE_COL)
    r = p.add_run(text)
    r.font.name = CODE_FONT
    r.font.size = Pt(size)
    return p


# bold before italic, so ** never gets eaten as two single-asterisk delimiters
INLINE = re.compile(r'(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`]+`|\[[^\]]+\]\([^)]+\))')


def emit(par, text):
    text = text.replace(r'\|', '|')
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            par.add_run(tok[1:-1]).italic = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = par.add_run(tok[1:-1])
            r.font.name = CODE_FONT
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
        elif tok.startswith('['):
            label, _url = re.match(r'\[([^\]]+)\]\(([^)]+)\)', tok).groups()
            r = par.add_run(label)
            r.font.color.rgb = RGBColor(0x1A, 0x5F, 0xB4)
            r.underline = True
        else:
            par.add_run(tok)


def split_row(line):
    return [c.strip() for c in re.split(r'(?<!\\)\|', line.strip())[1:-1]]


def new_doc():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = BODY_FONT
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(7)
    for i, sz in ((1, 19), (2, 14.5), (3, 12)):
        h = doc.styles[f'Heading {i}']
        h.font.name = BODY_FONT
        h.font.size = Pt(sz)
        h.font.color.rgb = INK
        h.font.bold = True
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.85)
        s.top_margin = s.bottom_margin = Inches(0.8)
    return doc


def render_markdown(doc, md):
    lines = md.split('\n')
    i = 0
    while i < len(lines):
        ln = lines[i]

        if ln.startswith('```'):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].startswith('```'):
                buf.append(lines[i])
                i += 1
            i += 1
            code_block(doc, '\n'.join(buf))
            continue

        if ln.strip().startswith('|') and i + 1 < len(lines) \
                and re.match(r'^\s*\|[\s:|-]+\|\s*$', lines[i + 1]):
            head = split_row(ln)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append(split_row(lines[i]))
                i += 1
            t = doc.add_table(rows=1, cols=len(head))
            t.style = 'Table Grid'
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for c, txt in zip(t.rows[0].cells, head):
                c.text = ''
                p = c.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                emit(p, txt)
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9.5)
                shade(c._tc.get_or_add_tcPr(), 'E9E6F2')
            for row in rows:
                cells = t.add_row().cells
                for c, txt in zip(cells, row + [''] * (len(head) - len(row))):
                    c.text = ''
                    p = c.paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    emit(p, txt)
                    for r in p.runs:
                        r.font.size = Pt(9)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        m = re.match(r'^(#{1,4})\s+(.*)$', ln)
        if m:
            lvl, txt = len(m.group(1)), m.group(2)
            p = doc.add_heading('', level=min(lvl, 3))
            p.paragraph_format.space_before = Pt(14 if lvl <= 2 else 10)
            p.paragraph_format.space_after = Pt(5)
            emit(p, txt)
            for r in p.runs:
                r.font.color.rgb = INK
            i += 1
            continue

        if re.match(r'^---+\s*$', ln):
            hrule(doc)
            i += 1
            continue

        m = re.match(r'^(\d+)\.\s+(.*)$', ln)
        if m:
            body = m.group(2)
            while i + 1 < len(lines) and lines[i + 1].startswith('   ') and lines[i + 1].strip():
                body += ' ' + lines[i + 1].strip()
                i += 1
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(4)
            emit(p, body)
            i += 1
            continue

        m = re.match(r'^(\s*)[-*]\s+(.*)$', ln)
        if m:
            body = m.group(2)
            while i + 1 < len(lines) and lines[i + 1].startswith('  ') and lines[i + 1].strip() \
                    and not re.match(r'^\s*[-*]\s', lines[i + 1]):
                body += ' ' + lines[i + 1].strip()
                i += 1
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            if len(m.group(1)) >= 2:
                p.paragraph_format.left_indent = Inches(0.6)
            emit(p, body)
            i += 1
            continue

        if not ln.strip():
            i += 1
            continue

        body = [ln.strip()]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
                r'^(#{1,4}\s|```|\||\s*[-*]\s|\d+\.\s|---+\s*$)', lines[i]):
            body.append(lines[i].strip())
            i += 1
        emit(doc.add_paragraph(), ' '.join(body))


def append_prompts(doc):
    """One page per vial, each holding that vial's complete prompt."""
    missing = [i for i in PRODUCTS if not (SEED / f'prompt-{i}.txt').exists()]
    if missing:
        raise SystemExit(f'missing prompt files for {missing} — '
                         f'run tools/emit_seedance_prompts.py first')

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    h = doc.add_heading('', level=1)
    emit(h, 'Appendix — the eight finished prompts')
    for r in h.runs:
        r.font.color.rgb = INK

    emit(doc.add_paragraph(),
         'Each block below is complete and paste-ready: select it, copy it, and '
         'drop the whole thing into the Seedance **Prompt** box. Nothing needs '
         'substituting. Upload that vial\'s `_seed/{id}.png` as **both** the '
         'image and the end frame before generating.')
    emit(doc.add_paragraph(),
         'These are generated from `assets/video/_seed/prompt-N.txt`. If you '
         'edit the prompt, edit the markdown source and re-run the two scripts '
         'rather than editing here — Word edits are overwritten on the next build.')

    for pid, name in PRODUCTS.items():
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        h = doc.add_heading('', level=2)
        emit(h, f'{pid}. {name}')
        for r in h.runs:
            r.font.color.rgb = INK
        text = (SEED / f'prompt-{pid}.txt').read_text(encoding='utf-8')
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(f'Start frame + end frame:  assets/video/_seed/{pid}.png'
                      f'          ({len(text)} characters, limit 5000)')
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x50, 0x66)
        r.italic = True
        code_block(doc, text, size=9)


if __name__ == '__main__':
    doc = new_doc()
    render_markdown(doc, SRC.read_text(encoding='utf-8'))
    append_prompts(doc)
    doc.save(DST)
    print(f'wrote {DST.relative_to(ROOT)}')
